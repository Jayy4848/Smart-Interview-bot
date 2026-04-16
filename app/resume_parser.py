from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def extract_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in {".txt", ".md"}:
        return _read_txt(path)
    raise ValueError(f"Unsupported resume file type: {suffix}")


def extract_resume_text_from_bytes(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(parts).strip()
    if suffix == ".docx":
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    if suffix in {".txt", ".md"}:
        return content.decode("utf-8", errors="ignore").strip()
    raise ValueError(f"Unsupported resume file type: {suffix}")

